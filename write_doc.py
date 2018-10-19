# !/usr/bin/python3
# -*- coding: utf-8 -*-
# @Time    : 2018-09-15 22:55
# @Author  : Jackadam
# @Email   :jackadam@sina.com
# @File    : write_doc.py
# @Software: PyCharm
import requests, json, sqlite3, time, random
import sqlalchemy
from sqlalchemy import create_engine
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy import Column, Integer, String,desc
from sqlalchemy.orm import sessionmaker
import docx
Base = declarative_base()  # 生成orm基类
engine = create_engine('sqlite:///Plan.db', encoding='utf-8')
Session_class = sessionmaker(bind=engine)  # 创建与数据库的会话session class ,注意,这里返回给session的是个class,不是实例
Session = Session_class()  # 生成session实例



class Question(Base):
    __tablename__ = 'Question'  # 表名
    id = Column(Integer, primary_key=True)
    Question_type = Column(String(32))
    Question = Column(String(128))
    Answer = Column(String(32))
    AnswerA = Column(String(32))
    AnswerB = Column(String(32))
    AnswerC = Column(String(32))
    AnswerD = Column(String(32))
    AnswerE = Column(String(32))
    AnswerF = Column(String(32))



def choice_Answer(cow,answer):
    if answer=='A':
        result='A：'+cow.AnswerA
    elif answer=='B':
        result='B：'+cow.AnswerB
    elif answer=='C':
        result='C：'+cow.AnswerC
    elif answer=='D':
        result='D：'+cow.AnswerD
    elif answer=='E':
        result='E：'+cow.AnswerE
    elif answer=='F':
        result='F：'+cow.AnswerF
    return result

def get_Answer(cow):
    result=[]
    if cow.Question_type=='判断题':
        result.append(cow.Answer)
    elif cow.Question_type=='单选题':
        result.append(choice_Answer(cow,cow.Answer))
    elif cow.Question_type=='多选题':
        for i in cow.Answer:
            result.append(choice_Answer(cow, i))
    return result


if __name__ == '__main__':
    test = file = docx.Document()


    info=Session.query(Question).order_by(Question.Question_type,Question.Question).all()
    lins=0
    for i in info:
        lins=lins+1
        test.add_paragraph('%s:(%s)  %s'%(lins,i.Question_type,i.Question))
        answer=get_Answer(i)
        str_answer='      答案：'
        for j in answer:
            str_answer=str_answer+j+'，'
        test.add_paragraph(str_answer)
    test.save('test.docx')